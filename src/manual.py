"""Yardım/manual sistemi — bir koleksiyon için çapraz-bağlantılı, çok bölümlü
teknik dokümantasyon üretir (HTML görüntüleme + PDF/DOCX dışa aktarma).

Tasarım: TEK bir belge modeli (düz dict — projenin geri kalanıyla aynı
sözleşme, retrieval.py'nin her yerde dict döndürmesiyle tutarlı) inşa edilip
kalıcı JSON olarak saklanır; HTML/PDF/DOCX bu TEK modelin üç ayrı
render'ıdır — mantık üçe katlanmaz. Zaten var olan yapı taşları üstüne kurulu:
  - document_unit(): dosya başına LLM özeti (ZATEN önbellekli — build_manual
    onu tekrar tekrar çağırsa bile ilk üretimden sonra bedava)
  - build_symbol_graph()/get_type_hierarchy(): tip adı -> hangi dosyada
    tanımlı bilgisi, HTML'deki çapraz-linkleme buradan gelir
  - get_profile_payload(): başlık sayfası (owner/group/kaynak/version)

Bölümleme: unit yolunun İLK klasör segmenti (örn. "jvcl/tests/..." ->
"jvcl") — büyük koleksiyonlarda (Jedi) tek düz sayfa yerine gezilebilir
kitap yapısı. `scope` verilirse yalnız o önekle başlayan dosyalar dahil
edilir (17K+ sembollü Jedi gibi bir koleksiyonun TAMAMI yerine bir alt-ağaç
için manual üretilebilsin diye).
"""
import io
import json
import pathlib
import re
import zipfile
from datetime import datetime, timezone
from urllib.parse import quote as _urlq

from qdrant_client import models

try:
    from . import retrieval
except ImportError:
    import retrieval

ROOT = pathlib.Path(__file__).resolve().parent.parent
MANUAL_DIR = ROOT / "data" / "manuals"

# Sıra 3/4 (kullanıcı, highlight.js seçildi): projenin chunker dil etiketleri
# (src/chunker.py LANG_TABLE) çoğunlukla highlight.js'in vendored dosya adlarıyla
# BİREBİR aynı — ama 3 bilinen istisna var (static/vendor/highlightjs/languages
# içeriğine göre doğrulandı): Pascal/Delphi dosyası "delphi.min.js" adında
# (kayıtlı takma adları: dpr/dfm/pas/pascal), Objective-C "objectivec.min.js",
# Visual Basic "vbnet.min.js". Burada YOKSA (ör. zig/solidity/gdscript/odin/
# gleam — hljs'te hiç yok) proje dil adı OLDUĞU GİBİ denenir; dosya 404 dönerse
# istemci tarafında SESSİZCE düz metne düşülür (bkz. script'teki manualHL) —
# bu yüzden liste TAM olmak ZORUNDA değil, yalnız BİLİNEN istisnaları düzeltir.
_HLJS_ALIAS = {"pascal": "delphi", "objc": "objectivec", "vb": "vbnet"}
_HLJS_UNSUPPORTED = {"zig", "solidity", "gdscript", "odin", "gleam"}


def _hljs_lang(lang: str) -> str:
    """Boş string dönerse istemci highlight.js'i hiç DENEMEZ (bilinen-desteksiz)."""
    if not lang or lang in _HLJS_UNSUPPORTED:
        return ""
    return _HLJS_ALIAS.get(lang, lang)


def _chapter_key(unit: str) -> str:
    parts = unit.replace("\\", "/").split("/")
    return parts[0] if len(parts) > 1 else "Genel"


def _slug(*parts: str) -> str:
    s = "-".join(parts).lower()
    s = re.sub(r"[^a-z0-9]+", "-", s).strip("-")
    return s or "bolum"


def _build_class_tree(sections: list[dict], chapter_types: list[dict], edges: list[dict]) -> dict:
    """Sıra B (kullanıcı, "Tree list Classlara göre yapılmalı iç içe"): bir
    bölümün DOSYA listesinin yanına sınıf hiyerarşisi ağacı ekler.
    build_symbol_graph()'ın ürettiği inherits/implements kenarlarından
    (retrieval.get_all_class_edges — TÜM koleksiyon için TEK seferde çekilir)
    bu bölüme ÖZEL, PURE (qdrant'a dokunmaz, test edilebilir) bir ağaç kurar.

    sections: bu bölümün TÜM dosyaları [{"unit","slug","title"}, ...]
              (chapters[ch]["sections"] ile aynı sözleşme — sınıfsız dosyalar
              dahil, "Diğer/Global" grubu için gerekir)
    chapter_types: bu bölümde TANIMLI sınıf/arayüz adları [{"name","unit"}, ...]
    edges: TÜM koleksiyonun kenarları (filtrelenmemiş) — {"child_name",
           "child_display","parent_name","parent_display","unit"}

    Döner: {"roots": [{"name","href","external","children":[...]}, ...],
            "other_files": [{"title","href"}, ...]}   # sınıfsız dosyalar

    Kullanıcı kararı: ebeveyni BU BÖLÜMDE tanımlı olmayan düğümler (harici RTL
    sınıfı — TObject/TComponent — YA DA başka bir bölümdeki sınıf, ikisi de
    burada AYIRT EDİLMEZ, bilinçli basitleştirme) "external": true, "href":
    None ile işaretlenir — arayüzde gri/tıklanamaz kök olarak gösterilir."""
    unit_to_href = {s["unit"]: f'#{s["slug"]}' for s in sections}
    name_to_unit = {t["name"].lower(): t["unit"] for t in chapter_types if t.get("name")}
    chapter_names = set(name_to_unit)

    children_of: dict[str, list[str]] = {}
    # KENARI OLMAYAN (izole, yalnız-kök) tipler için de doğru büyük/küçük harfli
    # ad gerekir — aşağıdaki döngü yalnız bir KENARDE geçen adları doldurur.
    display: dict[str, str] = {t["name"].lower(): t["name"] for t in chapter_types if t.get("name")}
    seen_child = set()
    for e in edges:
        cn = (e.get("child_name") or "").lower()
        if cn not in chapter_names:
            continue   # çocuk bu bölümde tanımlı değilse (başka bölüm) burada dal değil
        pn = (e.get("parent_name") or "").lower()
        if not pn:
            continue
        display[cn] = e.get("child_display") or cn
        display.setdefault(pn, e.get("parent_display") or pn)
        children_of.setdefault(pn, []).append(cn)
        seen_child.add(cn)

    def node(name: str, path: frozenset = frozenset()) -> dict:
        unit = name_to_unit.get(name)
        kids = [] if name in path else sorted(
            (node(c, path | {name}) for c in set(children_of.get(name, []))),
            key=lambda n: n["name"].lower())
        return {"name": display.get(name, name), "href": unit_to_href.get(unit) if unit else None,
                "external": unit is None, "children": kids}

    parent_names = set(children_of.keys())
    root_names = (parent_names - seen_child) | (chapter_names - seen_child - parent_names)
    roots = sorted((node(n) for n in root_names), key=lambda n: n["name"].lower())

    typed_units = set(name_to_unit.values())
    other_files = sorted(
        ({"title": s["title"], "href": f'#{s["slug"]}'} for s in sections if s["unit"] not in typed_units),
        key=lambda f: f["title"].lower())
    return {"roots": roots, "other_files": other_files}


def _topo_order(unit_paths: list[str], deps: dict[str, set[str]]) -> dict[str, int]:
    """Kahn algoritması — TEMEL (başkalarınca kullanılan/bağımlılığı az) dosyalar
    ÖNCE gelsin diye: `deps[u]` = u'nun (aynı bölüm içinde) bağımlı olduğu
    dosyalar. Bir dosya, TÜM bağımlılıkları zaten sıralanana kadar bekler.
    Çevrimli/çözülemeyen kalanlar en sona, alfabetik eklenir. Döner:
    {unit_path: sıra_indexi} — sections listesini bununla sort etmek için."""
    remaining = set(unit_paths)
    order: list[str] = []
    while remaining:
        ready = sorted(u for u in remaining if not (deps.get(u, set()) & remaining))
        if not ready:   # çevrim — kalanları alfabetik ekleyip çık
            order.extend(sorted(remaining))
            break
        order.extend(ready)
        remaining -= set(ready)
    return {u: i for i, u in enumerate(order)}


def build_manual(collection: str, scope: str = "", force: bool = False, st: dict | None = None, lang: str = "en") -> dict:
    """Koleksiyon için manual modelini kurar ve `data/manuals/<collection>.json`
    olarak kalıcı yazar. `force=True` document_unit önbelleğini de atlayıp
    tüm dosya özetlerini yeniden ürettirir (pahalı — normalde gerekmez, çünkü
    document_unit zaten kendi başına önbellekli).

    lang="en" (Sıra 5, kullanıcı kararı — "Sadece İngilizce yapalım, manuel
    içinde istediğimiz dile çevir gibi bir yer olsun"): manuel artık BAZ olarak
    İngilizce üretilir; Türkçe (veya başka bir dil) ayrı, isteğe bağlı bir AI
    ÇEVİRİ katmanı olarak translate_manual() ile SONRADAN eklenir, model.json'da
    saklanır — özgün İngilizce içerik asla kaybolmaz, çeviri onun ÜZERİNE
    yazılmaz."""
    cl = retrieval.cl
    prof = retrieval.get_profile_payload(collection)

    # ---- kapsamdaki dosyalar (unithead chunk'ları — hafif, tam kod gerekmiyor) ----
    units: list[dict] = []
    next_page = None
    while True:
        batch, next_page = cl.scroll(collection, limit=2000, offset=next_page,
            with_payload=["unit", "lang", "name", "uses"],
            scroll_filter=models.Filter(must=[models.FieldCondition(key="kind", match=models.MatchValue(value="unithead"))]))
        for p in batch:
            u = p.payload.get("unit") or ""
            if not scope or u.replace("\\", "/").startswith(scope):
                units.append({"unit": u, "lang": p.payload.get("lang") or "pascal", "chunk_id": p.id,
                              "name": p.payload.get("name") or "", "uses": p.payload.get("uses") or []})
        if next_page is None:
            break
    units.sort(key=lambda x: x["unit"])
    if not units:
        return {"error": f"kapsamda dosya bulunamadı (scope={scope!r}) — koleksiyon unithead içermiyor olabilir (v1 chunker ile indekslenmiş olabilir, yeniden indeksleyin)"}

    # ---- Sıra: "dosya sıralaması kullanım bağımlılıklarına göre olmalı" — TEMEL
    # (başkalarınca kullanılan) dosyalar önce, onları kullananlar sonra. `uses`
    # ADLARLA (unit adı/import string'i) geliyor — yalnız BU KAPSAMDAKİ diğer
    # unit'lerle eşleşenler gerçek kenar sayılır (dış bağımlılıklar sıralamayı
    # etkilemez, zaten çözülemez). Adlar hem TAM hem SON-NOKTALI-SEGMENT ile
    # eşleştirilir (Pascal "System.SysUtils" tam eşleşir; jenerik dillerde
    # import genelde son segmentle örtüşür, ör. "./utils" -> "utils.py")."""
    name_to_path: dict[str, str] = {}
    for u in units:
        if u["name"]:
            name_to_path.setdefault(u["name"].lower(), u["unit"])
        stem = pathlib.Path(u["unit"]).stem.lower()
        name_to_path.setdefault(stem, u["unit"])
    deps: dict[str, set[str]] = {}
    for u in units:
        resolved = set()
        for used in u["uses"]:
            used_l = str(used).strip().lower()
            target = name_to_path.get(used_l) or name_to_path.get(used_l.split(".")[-1]) or name_to_path.get(used_l.split("/")[-1])
            if target and target != u["unit"]:
                resolved.add(target)
        deps[u["unit"]] = resolved
    dep_order = _topo_order([u["unit"] for u in units], deps)

    if st is not None:
        st.update(phase="building", total=len(units), done=0)

    # ---- tip adı -> ev sahibi bölüm slug'ı (HTML çapraz-linkleme için) + Sıra B
    # (kullanıcı): aynı taramada bölüm başına tip listesi de toplanır — sınıf
    # ağacı (_build_class_tree) bunu kullanır, ayrı bir sorgu turu gerekmez.
    type_home: dict[str, str] = {}   # bare-lower isim -> "chapter-slug#section-slug"
    chapter_types: dict[str, list[dict]] = {}   # chapter_key -> [{"name","unit"}, ...]
    for u in units:
        ch = _chapter_key(u["unit"])
        sec = _slug(u["unit"])
        pts, _ = cl.scroll(collection, limit=200, with_payload=["name", "kind"],
            scroll_filter=models.Filter(must=[
                models.FieldCondition(key="unit", match=models.MatchValue(value=u["unit"])),
                models.FieldCondition(key="kind", match=models.MatchValue(value="type"))]))
        for p in pts:
            nm = (p.payload.get("name") or "").split("=")[0].strip()
            if nm:
                type_home[nm.lower()] = f"{_slug(ch)}.html#{sec}"
                chapter_types.setdefault(ch, []).append({"name": nm, "unit": u["unit"]})

    # ---- bölüm içeriği: document_unit'in ÖNBELLEKLİ özeti ----
    chapters: dict[str, dict] = {}
    lang_counts: dict[str, int] = {}
    for i, u in enumerate(units):
        if st is not None:
            st.update(done=i, current=u["unit"])
        lang_counts[u["lang"]] = lang_counts.get(u["lang"], 0) + 1
        doc = retrieval.document_unit(collection, u["unit"], force=force, lang=lang)
        if "error" in doc:
            continue
        ch_key = _chapter_key(u["unit"])
        chapters.setdefault(ch_key, {"title": ch_key, "slug": _slug(ch_key), "sections": []})
        chapters[ch_key]["sections"].append({
            "title": pathlib.Path(u["unit"]).name, "slug": _slug(u["unit"]),
            "unit": u["unit"], "lang": u["lang"], "body_md": doc["md"], "chunk_id": u["chunk_id"],
        })
    for ch in chapters.values():
        ch["sections"].sort(key=lambda s: dep_order.get(s["unit"], len(dep_order)))

    # ---- Sıra B (kullanıcı): sınıf-bazlı iç içe ağaç — TEK bulk kenar sorgusu,
    # sonra her bölüm için PURE _build_class_tree (qdrant'a bir daha dokunmaz).
    class_edges = retrieval.get_all_class_edges(collection)
    for ch_key, ch in chapters.items():
        ch["class_tree"] = _build_class_tree(ch["sections"], chapter_types.get(ch_key, []), class_edges)

    title = f"{collection} — Kullanım Kılavuzu" if lang == "tr" else f"{collection} — User Guide"
    model = {
        "collection": collection, "scope": scope, "lang": lang, "translations": {},
        "title": title,
        "owner": prof.get("owner", ""), "group": prof.get("group", ""),
        "version": prof.get("version", ""), "kaynak": prof.get("kaynak", ""), "url": prof.get("url", ""),
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "stats": {"units": len(units), "chapters": len(chapters), "languages": lang_counts},
        "chapters": sorted(chapters.values(), key=lambda c: c["title"].lower()),
        "type_home": type_home,
    }
    _save_manual(collection, model)
    if st is not None:
        st.update(done=len(units))
    return model


def load_manual(collection: str) -> dict | None:
    f = MANUAL_DIR / f"{collection}.json"
    if not f.exists():
        return None
    return json.loads(f.read_text(encoding="utf-8"))


def _save_manual(collection: str, model: dict):
    MANUAL_DIR.mkdir(parents=True, exist_ok=True)
    (MANUAL_DIR / f"{collection}.json").write_text(json.dumps(model, ensure_ascii=False, indent=1), encoding="utf-8")


LANG_LABELS = {"en": "English", "tr": "Türkçe"}   # bilinen birkaçı için düzgün ad — bilinmeyen kod .title() ile düşer


def model_for_lang(model: dict, lang: str) -> dict:
    """`model`in section body_md'lerini `lang` çevirisiyle DEĞİŞTİRİLMİŞ bir
    KOPYASINI döner (özgün model asla mutasyona uğramaz). lang boşsa veya baz
    dilse model AYNEN döner. HTML/DOCX/PDF render'larının HEPSİ bunu kullanır
    — çeviri mantığı üç render'da AYRI AYRI tekrarlanmaz, tek yerde."""
    base_lang = model.get("lang", "en")
    if not lang or lang == base_lang:
        return model
    translation = (model.get("translations") or {}).get(lang)
    if not translation:
        return model
    trans_secs = translation.get("sections") or {}
    out = dict(model)
    out["title"] = model["title"]   # başlık çevrilmiyor (kısa, koleksiyon adı içeriyor) — bilinçli
    new_chapters = []
    for ch in model["chapters"]:
        new_ch = dict(ch)
        new_ch["sections"] = [
            {**sec, "body_md": trans_secs.get(f'{ch["slug"]}|{sec["slug"]}', sec["body_md"])}
            for sec in ch["sections"]]
        new_chapters.append(new_ch)
    out["chapters"] = new_chapters
    return out


def list_manual_languages(collection: str) -> list[dict]:
    """Bu manuel için mevcut diller: baz (üretildiği dil) + eklenmiş çeviriler.
    UI'daki dil seçici bunu kullanır — "her eklenen dil orada görünsün" (Sıra 5)."""
    m = load_manual(collection)
    if m is None:
        return []
    base = m.get("lang", "en")
    out = [{"code": base, "label": LANG_LABELS.get(base, base.title()), "base": True}]
    for code, t in (m.get("translations") or {}).items():
        out.append({"code": code, "label": t.get("label") or LANG_LABELS.get(code, code.title()), "base": False})
    return out


def translate_manual(collection: str, target_lang: str, target_label: str = "", model: str = "",
                      force: bool = False, st: dict | None = None) -> dict:
    """Sıra 5 (kullanıcı): var olan (baz dilde, varsayılan İngilizce) manuel
    modelini AI ile hedef dile çevirir; sonucu model.json'un `translations`
    alanında KALICI saklar. Yapı (bölüm/section sırası, slug, chunk_id) AYNEN
    korunur — yalnız gövde metni (body_md) çevrilir. section BAŞINA önbelleklidir
    (force=False iken zaten çevrilmiş bir bölüm atlanır — force=True hepsini
    yeniden çevirir). Özgün baz-dil içeriği ASLA üzerine yazılmaz, translations
    ayrı bir alanda tutulur."""
    doc_model = load_manual(collection)
    if doc_model is None:
        return {"error": f"'{collection}' için manuel henüz üretilmemiş"}
    target_lang = (target_lang or "").strip().lower()
    if not target_lang:
        return {"error": "hedef dil boş olamaz"}
    base_lang = doc_model.get("lang", "en")
    if target_lang == base_lang:
        return {"error": f"manuel zaten '{target_lang}' dilinde üretilmiş — çeviriye gerek yok"}
    mdl = model or retrieval._CFG.get("deep_model", "qwen3.6")
    translations = doc_model.setdefault("translations", {})
    existing = {} if force else dict((translations.get(target_lang) or {}).get("sections") or {})
    total = sum(len(ch["sections"]) for ch in doc_model["chapters"])
    if st is not None:
        st.update(phase="translating", total=total, done=0)
    label = target_label.strip() if target_label and target_label.strip() else LANG_LABELS.get(target_lang, target_lang.title())
    done = 0
    for ch in doc_model["chapters"]:
        for sec in ch["sections"]:
            key = f'{ch["slug"]}|{sec["slug"]}'
            if st is not None:
                st.update(done=done, current=sec["unit"])
            done += 1
            if key in existing:
                continue
            prompt = (f"Translate the following technical documentation from English to {label}. "
                      "Preserve the EXACT Markdown structure (headings, bullet lists, code blocks) — do NOT "
                      "translate content inside code blocks (```...```) or inline code (`...`), and do NOT "
                      "translate identifier/type/function names. Output ONLY the translated Markdown text, "
                      "no extra commentary before or after.\n\n" + sec["body_md"])
            existing[key] = retrieval.ollama_generate(mdl, prompt, num_predict=1400)
    translations[target_lang] = {"label": label, "model": mdl,
                                  "generated_at": datetime.now(timezone.utc).isoformat(), "sections": existing}
    _save_manual(collection, doc_model)
    if st is not None:
        st.update(done=total)
    return {"ok": True, "lang": target_lang, "label": label, "sections": len(existing)}


# ---------------- HTML render (tek paylaşılan CSS, "hepsi aynı türden") ----------------
_LINK_SKIP_RE = re.compile(r"```.*?```|`[^`]*`", re.S)   # kod bloklarına/satır-içi koda dokunma

def _cross_link(body_md: str, type_home: dict[str, str], self_href: str) -> str:
    """Bilinen tip adlarını [Ad](hedef) markdown linkine çevirir — kod
    bloklarını ve satır-içi kodu atlar, kendi sayfasına link vermez."""
    if not type_home:
        return body_md
    names = sorted(type_home, key=len, reverse=True)   # uzun adlar önce (alt-string çakışmasını önler)
    pattern = re.compile(r"\b(" + "|".join(re.escape(n) for n in names) + r")\b", re.I)

    def repl_outside_code(segment: str) -> str:
        def repl(m):
            href = type_home.get(m.group(1).lower())
            if not href or href == self_href:
                return m.group(1)
            return f"[{m.group(1)}]({href})"
        return pattern.sub(repl, segment)

    out, last = [], 0
    for m in _LINK_SKIP_RE.finditer(body_md):
        out.append(repl_outside_code(body_md[last:m.start()]))
        out.append(m.group(0))   # kod bloğu/satır-içi kod OLDUĞU GİBİ
        last = m.end()
    out.append(repl_outside_code(body_md[last:]))
    return "".join(out)


def _md_to_html_fragment(md: str, hljs_lang: str = "") -> str:
    """viewer.html'deki mdToHtml'in Python karşılığı — aynı minimal alt küme
    (başlık/kalın/italik/kod/link/liste), tam CommonMark değil, bilinçli.

    Sıra 5 (kullanıcı): kod gösteriminde her yerde highlight.js kullanılsın —
    body_md içindeki ```kod bloğu``` parçaları burada da var (ör. "Public API"
    örnekleri); document_unit() tüm body_md'yi TEK bir dosya/dil için ürettiğinden
    (bkz. build_manual) o bölümün `sec["lang"]`ı BURADAKİ TÜM kod bloklarına da
    uygulanır — ayrı ayrı dil tespiti gerekmez."""
    def esc(s): return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    blocks = []
    code_blocks = []
    def stash_code(m):
        code_blocks.append(m.group(1))
        return f"\x00CODE{len(code_blocks) - 1}\x00"
    md = re.sub(r"```\w*\n?(.*?)```", stash_code, md, flags=re.S)

    def inline(t):
        t = re.sub(r"`([^`]+)`", lambda m: f"<code>{esc(m.group(1))}</code>", t)
        t = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", t)
        t = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', t)
        return t

    html, in_list = [], False
    for line in md.split("\n"):
        m = re.match(r"^(#{1,3})\s+(.*)", line)
        if m:
            if in_list: html.append("</ul>"); in_list = False
            lvl = len(m.group(1))
            html.append(f"<h{lvl}>{inline(esc(m.group(2)))}</h{lvl}>")
            continue
        m = re.match(r"^[-*]\s+(.*)", line)
        if m:
            if not in_list: html.append("<ul>"); in_list = True
            html.append(f"<li>{inline(esc(m.group(1)))}</li>")
            continue
        if in_list: html.append("</ul>"); in_list = False
        if line.strip():
            html.append(f"<p>{inline(esc(line))}</p>")
    if in_list: html.append("</ul>")
    out = "\n".join(html)
    code_cls = f' class="language-{hljs_lang}"' if hljs_lang else ""
    for i, code in enumerate(code_blocks):
        out = out.replace(f"\x00CODE{i}\x00", f"<pre><code{code_cls}>{esc(code)}</code></pre>")
    return out


MANUAL_CSS = """
:root{
  --bg:#14110c;--panel:#1b1710;--card:#211c14;--code:#0f0d09;
  --line:#332b1f;--line2:#463b2a;--txt:#f1ebdd;--dim:#a99a83;--faint:#6f6555;
  --amber:#e0a24a;--amber-d:#c58a37;--teal:#49b39c;--teal-d:#2f7e6c;--err:#e0704a;
  --serif:Georgia,'Iowan Old Style','Palatino Linotype',serif;
  --mono:'Cascadia Code','JetBrains Mono',Consolas,monospace;
}
*{box-sizing:border-box;margin:0}
body{background:var(--bg);color:var(--txt);font:15px/1.7 'Segoe UI',system-ui,sans-serif;display:flex;min-height:100vh}
nav{width:300px;flex:none;background:var(--panel);border-right:1px solid var(--line);padding:22px 18px;overflow-y:auto;overflow-x:hidden;position:sticky;top:0;height:100vh}
#navresize{position:absolute;right:-4px;top:0;width:8px;height:100%;cursor:ew-resize;z-index:5}
#navresize:hover,#navresize.dragging{background:rgba(224,162,74,.18)}
nav.dragging{transition:none}
nav h1{font-family:var(--serif);font-size:18px;color:var(--amber);margin-bottom:4px;line-height:1.3}
nav h1 a{color:inherit;text-decoration:none}
nav h1 a:hover{color:var(--teal)}
nav .home{display:inline-block;font-size:11.5px;color:var(--faint);text-decoration:none;margin-bottom:10px}
nav .home:hover{color:var(--amber)}
nav .meta{font-size:11.5px;color:var(--faint);margin-bottom:16px;font-family:var(--mono)}
nav input{width:100%;background:var(--card);border:1px solid var(--line2);border-radius:8px;color:var(--txt);padding:8px 10px;font-size:13px;margin-bottom:14px}
nav .chapter{margin-bottom:4px}
nav .chapter>a{display:block;font-size:12.5px;letter-spacing:.2px;color:var(--dim);padding:6px 4px;text-decoration:none;font-weight:700;overflow-wrap:anywhere}
nav .chapter a.sec{display:block;font-size:13px;color:var(--txt);padding:4px 4px 4px 14px;text-decoration:none;border-left:2px solid transparent;overflow-wrap:anywhere;line-height:1.35}
nav .chapter a.sec:hover{color:var(--amber);border-left-color:var(--amber-d)}
nav a.active{color:var(--amber)!important;border-left-color:var(--amber)!important}
main{flex:1;max-width:840px;margin:0 auto;padding:48px 40px 100px}
main h1,main h2,main h3{font-family:var(--serif);color:var(--amber);text-wrap:balance;margin:28px 0 12px;line-height:1.3}
main h1{font-size:30px;border-bottom:1px solid var(--line);padding-bottom:14px}
main h2{font-size:21px}
main h3{font-size:16px;color:var(--teal)}
main p{margin:0 0 14px;max-width:70ch}
main ul{margin:0 0 14px 22px}
main li{margin-bottom:4px}
main a{color:var(--teal)}
main code{background:var(--card);border:1px solid var(--line);border-radius:4px;padding:1px 6px;font-family:var(--mono);font-size:13px}
main pre{background:var(--code);border:1px solid var(--line);border-radius:10px;padding:14px 16px;overflow-x:auto;margin:0 0 16px;font:12.5px/1.6 var(--mono)}
main pre code{background:none;border:0;padding:0}
.section{padding-top:8px;border-top:1px solid var(--line);margin-top:32px}
.section:first-of-type{border-top:0;margin-top:0}
.badge{display:inline-block;font-size:10.5px;color:var(--teal);background:rgba(73,179,156,.1);padding:2px 8px;border-radius:6px;font-family:var(--mono);margin-left:8px;vertical-align:middle}
.overview-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(140px,1fr));gap:12px;margin:20px 0 28px}
.stat{background:var(--card);border:1px solid var(--line);border-radius:10px;padding:14px 16px}
.stat b{display:block;font-size:22px;color:var(--amber);font-family:var(--serif)}
.stat span{font-size:11.5px;color:var(--faint);text-transform:uppercase;letter-spacing:.5px}
.sec-actions{display:inline-flex;flex-wrap:wrap;gap:6px;margin-left:10px;vertical-align:middle}
.sec-actions button{font-size:11px;padding:3px 9px;background:var(--card);border:1px solid var(--line2);color:var(--dim);border-radius:6px;cursor:pointer;font-family:inherit}
.sec-actions button:hover{border-color:var(--amber-d);color:var(--amber)}
.sec-actions button:disabled{opacity:.5;cursor:wait}
#sidepanel{position:fixed;top:0;right:0;width:min(640px,92vw);height:100%;background:var(--panel);border-left:1px solid var(--line2);box-shadow:-12px 0 32px rgba(0,0,0,.4);z-index:30;transform:translateX(100%);transition:transform .25s ease;display:flex;flex-direction:column}
#sidepanel.show{transform:translateX(0)}
#sidepanel.dragging{transition:none}
#sp-resize{position:absolute;left:-4px;top:0;width:8px;height:100%;cursor:ew-resize;z-index:31}
#sp-resize:hover,#sp-resize.dragging{background:rgba(224,162,74,.18)}
.sp-head{display:flex;align-items:center;gap:10px;padding:14px 18px;border-bottom:1px solid var(--line);background:var(--card)}
.sp-fname{font-family:var(--mono);font-size:13px;color:var(--amber);overflow:hidden;text-overflow:ellipsis;white-space:nowrap;flex:1}
.sp-close{background:none;border:0;color:var(--faint);font-size:16px;cursor:pointer;padding:2px 8px}
.sp-close:hover{color:var(--amber)}
.sp-body{flex:1;overflow:auto;padding:0}
.sp-body pre{margin:0;padding:16px 20px;font:12.5px/1.6 var(--mono);color:#d8cbb0;white-space:pre;overflow-x:auto}
.hljs{background:transparent;color:var(--txt)}
.hljs-comment,.hljs-quote{color:var(--faint);font-style:italic}
.hljs-keyword,.hljs-selector-tag,.hljs-literal,.hljs-subst,.hljs-name,.hljs-tag{color:var(--amber)}
.hljs-string,.hljs-doctag,.hljs-regexp,.hljs-addition{color:var(--teal)}
.hljs-number,.hljs-symbol,.hljs-bullet,.hljs-link{color:var(--amber-d)}
.hljs-title,.hljs-type,.hljs-class .hljs-title,.hljs-built_in,.hljs-attr,.hljs-attribute{color:var(--teal-d)}
.hljs-variable,.hljs-template-variable,.hljs-params{color:var(--dim)}
.hljs-deletion{color:var(--err)}
.hljs-emphasis{font-style:italic}
.hljs-strong{font-weight:700}
.sp-err{padding:20px;color:var(--err)}
.spin{display:inline-block;width:13px;height:13px;border:2px solid var(--line2);border-top-color:var(--amber);border-radius:50%;animation:sp .7s linear infinite;vertical-align:-2px;margin-right:7px}
@keyframes sp{to{transform:rotate(360deg)}}
.langsw{display:flex;flex-wrap:wrap;gap:6px;margin-bottom:12px}
.langsw a{font-size:11.5px;padding:3px 9px;border-radius:99px;border:1px solid var(--line2);color:var(--dim);text-decoration:none}
.langsw a:hover{border-color:var(--amber-d);color:var(--amber)}
.langsw a.cur{background:var(--amber);border-color:var(--amber);color:#1a1305!important;font-weight:700}
.langsw a.addlang{color:var(--teal);border-style:dashed}
.langsw a.addlang:hover{border-color:var(--teal-d);color:var(--teal-d)}
.exportmenu{display:flex;flex-wrap:wrap;gap:6px;margin-bottom:14px}
.exportmenu a{font-size:11.5px;padding:3px 9px;border-radius:99px;border:1px solid var(--line2);color:var(--dim);text-decoration:none}
.exportmenu a:hover{border-color:var(--teal-d);color:var(--teal)}
.swal2-popup.ci-swal{background:var(--panel);color:var(--txt);border:1px solid var(--line2);border-radius:14px;font:15px/1.6 'Segoe UI',system-ui,sans-serif}
.ci-swal .swal2-title{color:var(--txt);font-size:18px}
.ci-swal .swal2-html-container{color:var(--dim)}
.ci-swal .swal2-input{background:var(--card);border:1px solid var(--line2);color:var(--txt);box-shadow:none;font-size:14px}
.ci-swal .swal2-input:focus{border-color:var(--amber-d);box-shadow:0 0 0 3px rgba(224,162,74,.08)}
.ci-swal-btn{background:var(--amber)!important;color:#1a1305!important;border:0!important;box-shadow:none!important;font-weight:700!important;border-radius:10px!important}
.ci-swal-btn:hover{background:var(--amber-d)!important}
.ci-swal-btn-cancel{background:var(--card)!important;color:var(--dim)!important;border:1px solid var(--line2)!important;box-shadow:none!important;border-radius:10px!important}
.ci-swal-toast{background:var(--panel)!important;color:var(--txt)!important;border:1px solid var(--line2);box-shadow:0 8px 24px rgba(0,0,0,.35)}
.ci-swal-toast .swal2-title{font-size:14px;color:var(--txt)}
.navtabs{display:flex;gap:4px;margin-bottom:10px}
.navtabs button{flex:1;font-size:11.5px;padding:5px 8px;background:var(--card);border:1px solid var(--line2);color:var(--dim);border-radius:6px;cursor:pointer;font-family:inherit}
.navtabs button.cur{background:var(--amber);border-color:var(--amber);color:#1a1305;font-weight:700}
.navview{display:none}
.navview.cur{display:block}
.ctree{list-style:none;margin:0;padding-left:0}
.ctree ul{list-style:none;margin:2px 0 2px 14px;padding-left:10px;border-left:1px dashed var(--line2)}
.ctree li{margin:2px 0}
.ctree a,.ctree span.ext{display:block;font-size:12.5px;padding:3px 4px;text-decoration:none;overflow-wrap:anywhere;line-height:1.35}
.ctree a{color:var(--txt)}
.ctree a:hover{color:var(--amber)}
.ctree span.ext{color:var(--faint);font-style:italic}
.ctree .other{margin-top:10px;padding-top:8px;border-top:1px solid var(--line)}
.ctree .other-label{font-size:11px;color:var(--faint);text-transform:uppercase;letter-spacing:.4px;margin-bottom:4px;padding-left:4px}
/* Sıra B/2. tur (kullanıcı): sınıf ağacı açılır/kapanır — native <details>
   kullanılıyor (JS/max-height transition YOK — bu oturumda max-height geçişi
   iki kez donma hatası verdi, bkz. Sıra 5/7; <details> tarayıcı-yerli, JS'siz
   ve o hata sınıfından bağışık). Varsayılan: kapalı; yalnız o an görüntülenen
   bölümün ağacı açık gelir ("auto_open" — render zamanında hesaplanır).*/
.ctree details{margin:0}
.ctree summary{cursor:pointer;list-style:none;display:flex;align-items:center;gap:2px}
.ctree summary::-webkit-details-marker{display:none}
.ctree summary::before{content:'▸';color:var(--faint);font-size:9px;flex:none;width:12px;transition:transform .15s ease}
.ctree details[open]>summary::before{transform:rotate(90deg)}
.ctree summary a,.ctree summary span.ext{flex:1;min-width:0;padding-left:0}
.ctree li>a,.ctree li>span.ext{padding-left:16px}
"""


def _render_ctree_node(n: dict, ch_slug: str, base: str, qs: str, auto_open: bool = False) -> str:
    """_build_class_tree'nin döndürdüğü TEK bir düğümü (ve alt ağacını) HTML'e
    çevirir. "external" düğümler (kullanıcı kararı: gri/tıklanamaz kök) `<span>`
    olarak, kendi bölümü olanlar `<a>` olarak render edilir.

    Sıra B/2. tur (kullanıcı): "Sınıf Ağacını Açılır Kapanır TreeList yap" —
    çocuğu olan düğümler <details>/<summary> ile (JS'siz, tarayıcı-yerli)
    aç/kapa olur; yapraklar (çocuksuz) düz link/span kalır. `auto_open`
    (çağıran taraf: `page == ch["slug"]`) o an görüntülenen bölümün ağacını
    varsayılan AÇIK, diğer bölümlerinkini KAPALI başlatır."""
    if n["href"]:
        label = f'<a href="{base}{ch_slug}.html{qs}{n["href"]}">{_esc(n["name"])}</a>'
    else:
        label = f'<span class="ext" title="Bu koleksiyonda tanımlı değil (harici sınıf veya başka bir bölümde)">{_esc(n["name"])}</span>'
    if not n["children"]:
        return label
    kids = "".join(f'<li>{_render_ctree_node(c, ch_slug, base, qs, auto_open)}</li>' for c in n["children"])
    open_attr = " open" if auto_open else ""
    return f'<details{open_attr}><summary>{label}</summary><ul>{kids}</ul></details>'


def render_manual_html(model: dict, page: str = "index", lang: str = "") -> str:
    """page: "index" (kapak+TOC) veya bir chapter slug'ı. lang: boşsa BAZ dil
    (model["lang"]); doluysa model["translations"][lang] varsa o dilin
    body_md'si kullanılır (bir bölüm henüz çevrilmemişse baz dile SESSİZCE düşer
    — hiç eksik/boş görünmez). Sıra 5 (kullanıcı): "manuel içinde istediğimiz
    dile çevir gibi bir yer olsun" — nav'daki dil seçici bunu sağlar.

    KRİTİK düzeltme: tüm href'ler MUTLAK yol (`/manual/{collection}/...`) olmalı.
    Eskiden `"{slug}.html"` gibi GÖRECELİ üretiliyordu — bu yalnız sayfa URL'si
    SONUNDA `/` varken doğru çözülür; ama gerçek sayfa `/manual/{collection}`
    (eğik çizgisiz) adresinde sunuluyor, bu yüzden tarayıcı `src.html`'i
    `/manual/src.html` diye çözüyordu (son segment "RESTRequest4Delphi"
    DEĞİŞTİRİLİYOR, ALTINA eklenmiyor) — canlı doğrulamada yakalanan gerçek
    hata: HER link "Manual henüz üretilmemiş" gösteriyordu. Var olan
    manual.json dosyaları (type_home içinde hâlâ göreli yol saklıyor)
    YENİDEN ÜRETİLMEDEN çalışsın diye düzeltme RENDER anında yapılıyor."""
    base_lang = model.get("lang", "en")
    active_lang = lang or base_lang
    qs = f"?lang={active_lang}" if active_lang != base_lang else ""
    display_model = model_for_lang(model, active_lang)   # body_md'ler çeviriyle değiştirilmiş kopya (varsa)

    base = f"/manual/{model['collection']}/"
    index_url = f"/manual/{model['collection']}{qs}"
    # list_manual_languages(collection) DİSKTEN okur — burada BİLEREK kullanılmıyor:
    # bu fonksiyon her zaman `model` PARAMETRESİYLE tutarlı olmalı (model_for_lang de
    # aynı sözleşmeyi izliyor), aksi halde diskteki hal ile bellekteki `model` farklıysa
    # (ör. henüz kaydedilmemiş bir model, testler) dil rozeti YANLIŞ/eksik listelenirdi.
    langs_avail = [{"code": base_lang, "label": LANG_LABELS.get(base_lang, base_lang.title()), "base": True}] + [
        {"code": code, "label": t.get("label") or LANG_LABELS.get(code, code.title()), "base": False}
        for code, t in (model.get("translations") or {}).items()]
    # class="cur" (NOT "active"): "active" çakışıyordu — nav a.active{color:var(--amber)
    # !important} kuralı (chapter/section vurgusu için) genel "nav a.active" seçicisiyle
    # dil rozetini de eşliyordu, !important yüzünden .langsw a.active'i eziyor, amber
    # ÜSTÜNE amber (görünmez metin) çıkıyordu — canlı doğrulamada yakalanan gerçek hata.
    lang_switcher = '<div class="langsw">' + "".join(
        f'<a href="{"/manual/" + model["collection"] + ("/" + page + ".html" if page != "index" else "")}'
        f'{"?lang=" + l["code"] if l["code"] != base_lang else ""}" '
        f'class="{"cur" if l["code"] == active_lang else ""}">{_esc(l["label"])}</a>'
        for l in langs_avail) + (
        '<a href="#" class="addlang" onclick="manualAddLang(event)" title="Yapay zeka ile yeni bir dile çevir">+ Dil ekle</a>'
    ) + '</div><div id="langjob" class="note" style="display:none"></div>'
    # Sıra 9 (kullanıcı): "dil menüsünün altına Export menüsü eklensin" — PDF/DOCX
    # zaten çalışan uçları + Sıra 8'in statik ZIP ucu, hepsi lang= parametresiyle
    # şu an görüntülenen dili dışa aktarır (baz dilse lang parametresi eklenmez).
    export_qs = f"&lang={_urlq(active_lang)}" if active_lang != base_lang else ""
    coll_q = _urlq(model["collection"])
    export_menu = ('<div class="exportmenu">'
        f'<a href="/api/manual/export?collection={coll_q}&format=pdf{export_qs}" title="PDF olarak indir">📄 PDF</a>'
        f'<a href="/api/manual/export?collection={coll_q}&format=docx{export_qs}" title="Word (DOCX) olarak indir">📝 DOCX</a>'
        f'<a href="/api/manual/export?collection={coll_q}&format=zip{export_qs}" title="Sunucu gerekmeden dosyada gezinilebilen statik HTML paketi (ZIP)">🗂️ Statik HTML (ZIP)</a>'
        '</div>')
    nav_html = ([f'<a class="home" href="{index_url}">← İndeks / Kapak</a>'] if page != "index" else []) + [
                f'<h1><a href="{index_url}">{_esc(model["title"])}</a></h1>',
                f'<div class="meta">{_esc(model.get("version",""))} '
                f'{"· " + _esc(model["owner"]) if model.get("owner") else ""}</div>',
                lang_switcher,
                export_menu]

    files_view = []
    for ch in model["chapters"]:
        active_ch = " active" if page == ch["slug"] else ""
        files_view.append(f'<div class="chapter"><a href="{base}{ch["slug"]}.html{qs}" class="{active_ch}">{_esc(ch["title"])}</a>')
        for sec in ch["sections"]:
            files_view.append(f'<a class="sec" href="{base}{ch["slug"]}.html{qs}#{sec["slug"]}">{_esc(sec["title"])}</a>')
        files_view.append("</div>")

    # Sıra B (kullanıcı, "Tree list Classlara göre yapılmalı iç içe"): dosya
    # listesinin yanına sekmeli bir sınıf-hiyerarşisi görünümü — yalnız
    # class_tree'si olan (build_manual Sıra B'yi destekleyen bir sürümle
    # üretilmiş) koleksiyonlarda gösterilir; eski manual.json'larda hiç
    # görünmez (geriye uyumlu, "Sınıflar" sekmesi hiç render edilmez).
    class_view = []
    has_any_classes = False
    for ch in model["chapters"]:
        ct = ch.get("class_tree") or {"roots": [], "other_files": []}
        if not ct["roots"] and not ct["other_files"]:
            continue
        has_any_classes = has_any_classes or bool(ct["roots"])
        active_ch = " active" if page == ch["slug"] else ""
        class_view.append(f'<div class="chapter"><a href="{base}{ch["slug"]}.html{qs}" class="{active_ch}">{_esc(ch["title"])}</a>')
        if ct["roots"]:
            items = "".join(f'<li>{_render_ctree_node(n, ch["slug"], base, qs, page == ch["slug"])}</li>' for n in ct["roots"])
            class_view.append(f'<ul class="ctree">{items}</ul>')
        if ct["other_files"]:
            of = "".join(f'<li><a href="{base}{ch["slug"]}.html{qs}{f["href"]}">{_esc(f["title"])}</a></li>' for f in ct["other_files"])
            class_view.append(f'<div class="other"><div class="other-label">Diğer / Global</div><ul class="ctree">{of}</ul></div>')
        class_view.append("</div>")

    if has_any_classes:
        nav_html.append('<div class="navtabs">'
            '<button type="button" class="cur" data-tab="files" onclick="manualNavTab(this)">Dosyalar</button>'
            '<button type="button" data-tab="classes" onclick="manualNavTab(this)">Sınıflar</button>'
            '</div>')
        nav_html.append('<input placeholder="Ara…" onkeyup="filterNav(this.value)" id="navsearch">')
        nav_html.append(f'<div class="navview cur" id="navview-files">{"".join(files_view)}</div>')
        nav_html.append(f'<div class="navview" id="navview-classes">{"".join(class_view)}</div>')
    else:
        nav_html.append('<input placeholder="Ara…" onkeyup="filterNav(this.value)" id="navsearch">')
        nav_html.extend(files_view)
    nav = "\n".join(nav_html)

    if page == "index":
        pl = ", ".join(f"{k} ({v})" for k, v in sorted(model["stats"]["languages"].items(), key=lambda kv: -kv[1]))
        body = f"""<h1>{_esc(model["title"])}</h1>
<p>{_esc(model.get("group","") or "")}</p>
<div class="overview-grid">
  <div class="stat"><b>{model["stats"]["units"]}</b><span>Dosya</span></div>
  <div class="stat"><b>{model["stats"]["chapters"]}</b><span>Bölüm</span></div>
</div>
<p><b>Diller:</b> {_esc(pl)}</p>
{"<p><b>Kaynak:</b> " + _esc(model.get("kaynak","")) + (" — <a href=\"" + _esc(model.get("url","")) + "\">" + _esc(model.get("url","")) + "</a>" if model.get("url") else "") + "</p>" if model.get("kaynak") else ""}
<h2>İçindekiler</h2>
<ul>{"".join(f'<li><a href="{base}{c["slug"]}.html{qs}">{_esc(c["title"])}</a> ({len(c["sections"])})</li>' for c in model["chapters"])}</ul>"""
    else:
        ch = next((c for c in display_model["chapters"] if c["slug"] == page), None)
        if ch is None:
            return "<h1>404</h1>"
        # BUG (kullanıcı, canlı doğrulamada bulundu): `v` zaten "{slug}.html#{section}"
        # biçiminde (FRAGMENT dahil) — qs SONA eklenirse "...html#section?lang=tr" çıkıyordu;
        # bu GEÇERSİZ bir URL'dir, çünkü "?query" HER ZAMAN "#fragment"TAN ÖNCE gelmeli.
        # Tarayıcı bunu "...html" + fragment="section?lang=tr" diye yorumluyordu — sonuç:
        # dil parametresi tamamen KAYBOLUYOR (hedef sayfa baz/İngilizce dilde açılıyordu) VE
        # kaydırma da çalışmıyordu (öyle bir id yok). qs artık "#" DAN ÖNCE ekleniyor.
        type_home_abs = {k: base + (v.replace("#", qs + "#", 1) if qs else v) for k, v in model.get("type_home", {}).items()}
        parts = [f'<h1>{_esc(ch["title"])}</h1>']
        for sec in ch["sections"]:
            linked = _cross_link(sec["body_md"], type_home_abs, f'{base}{ch["slug"]}.html{qs}#{sec["slug"]}')
            # "Aç" (kayıtlı varsayılan uygulamada) + "Tarayıcıda Göster" (sayfa içi
            # aç-kapa, gerçek kaynak) — YALNIZ chunk_id varsa (eski manual.json'lar
            # bu alanı henüz taşımıyor olabilir, yeniden üretilmeden de çökmesin).
            actions = ""
            if sec.get("chunk_id") is not None:
                cid = sec["chunk_id"]
                coll_esc = _esc(model["collection"])
                actions = (f'<span class="sec-actions">'
                          f'<button type="button" onclick="manualOpen(\'{coll_esc}\',{cid},this)" title="Kayıtlı varsayılan uygulamada aç">📂 Aç</button>'
                          f'<button type="button" onclick="manualOpenFolder(\'{coll_esc}\',{cid},this)" title="Dosya gezgininde, seçili olarak göster">🗂️ Klasörde Aç</button>'
                          f'<button type="button" onclick="manualShowInBrowser(\'{coll_esc}\',{cid},\'{_esc(sec["title"])}\',this,\'{_hljs_lang(sec["lang"])}\')">🖥️ Tarayıcıda Göster</button>'
                          f'</span>')
            # madde 2. tur, 2 (kullanıcı): "browserde aç" bölümün İÇİNE gömülü render
            # ediyordu, mantıksız — artık index.html'deki #sidepanel ile AYNI sağdan
            # kayan panelde açılıyor (bkz. manualShowInBrowser); bölüm içine gömülü
            # .sec-code div'i KALDIRILDI.
            parts.append(f'<div class="section" id="{sec["slug"]}"><h2>{_esc(sec["title"])}'
                         f'<span class="badge">{_esc(sec["lang"])}</span>{actions}</h2>'
                         f'{_md_to_html_fragment(linked, _hljs_lang(sec["lang"]))}</div>')
        body = "\n".join(parts)

    html_lang = active_lang if active_lang in ("tr", "en") else "en"
    return f"""<!doctype html><html lang="{html_lang}"><head><meta charset="utf-8">
<title>{_esc(model["title"])}</title><style>{MANUAL_CSS}</style>
<script src="/static/vendor/sweetalert2.min.js"></script>
<script src="/static/vendor/highlightjs/highlight.min.js"></script></head>
<body><nav id="manualnav">{nav}<div id="navresize"></div></nav><main>{body}</main>
<div id="sidepanel">
  <div id="sp-resize"></div>
  <div class="sp-head">
    <span class="sp-fname" id="sp-fname">…</span>
    <button class="sp-close" onclick="copySidePanel()" id="sp-copy" title="Kopyala">📋</button>
    <button class="sp-close" onclick="closeSidePanel()">✕</button>
  </div>
  <div class="sp-body" id="sp-body"></div>
</div>
<script>
var MANUAL_COLLECTION={json.dumps(model["collection"])}, MANUAL_PAGE={json.dumps(page)};
(function(){{
  var nav=document.getElementById('manualnav'),handle=document.getElementById('navresize');
  var saved=parseInt(localStorage.getItem('ci-manual-navw'));
  if(saved)nav.style.width=saved+'px';
  var dragging=false;
  handle.addEventListener('mousedown',function(e){{dragging=true;nav.classList.add('dragging');handle.classList.add('dragging');e.preventDefault();}});
  window.addEventListener('mousemove',function(e){{
    if(!dragging)return;
    var w=Math.min(window.innerWidth*0.6,Math.max(200,e.clientX));
    nav.style.width=w+'px';
  }});
  window.addEventListener('mouseup',function(){{
    if(!dragging)return;
    dragging=false;nav.classList.remove('dragging');handle.classList.remove('dragging');
    localStorage.setItem('ci-manual-navw',parseInt(nav.style.width));
  }});
}})();
// madde 2. tur, 2 (kullanıcı): "browserde aç" sağdan kayan panelde açılmalı —
// index.html'deki #sidepanel/#sp-resize deseniyle AYNI (sürükle-genişlet dahil).
var SP_CONTENT='';
// madde 2. tur, 3/4/5 (kullanıcı): highlight.js self-hosted — dil dosyası YALNIZ
// gerektiğinde (kaynak gösterildiğinde) dinamik <script> ile yüklenir, tekrar
// yüklenmez (HLJS_LOADED). Dosya yoksa (bkz. _hljs_lang — birkaç niş dil hljs'te
// yok) onerror SESSİZCE düz metne düşer, kullanıcı hata görmez.
var HLJS_LOADED={{}};
function manualHL(codeEl,hljsLang){{
  if(!hljsLang||typeof hljs==='undefined')return;
  codeEl.className='language-'+hljsLang;
  if(HLJS_LOADED[hljsLang]==='ok'){{hljs.highlightElement(codeEl);return;}}
  if(HLJS_LOADED[hljsLang]==='error')return;
  var s=document.createElement('script');
  s.src='/static/vendor/highlightjs/languages/'+hljsLang+'.min.js';
  s.onload=function(){{HLJS_LOADED[hljsLang]='ok';hljs.highlightElement(codeEl);}};
  s.onerror=function(){{HLJS_LOADED[hljsLang]='error';}};
  document.head.appendChild(s);
}}
function closeSidePanel(){{document.getElementById('sidepanel').classList.remove('show');}}
async function copySidePanel(){{
  if(!SP_CONTENT)return;
  try{{
    await navigator.clipboard.writeText(SP_CONTENT);
    var b=document.getElementById('sp-copy'),old=b.textContent;b.textContent='✅';
    setTimeout(function(){{b.textContent=old;}},1200);
  }}catch(e){{}}
}}
document.addEventListener('keydown',function(e){{
  if(e.key==='Escape'&&document.getElementById('sidepanel').classList.contains('show'))closeSidePanel();
}});
(function(){{
  var panel=document.getElementById('sidepanel'),handle=document.getElementById('sp-resize');
  var saved=parseInt(localStorage.getItem('ci-manual-spwidth'));
  if(saved)panel.style.width=saved+'px';
  var dragging=false;
  handle.addEventListener('mousedown',function(e){{dragging=true;panel.classList.add('dragging');handle.classList.add('dragging');e.preventDefault();}});
  window.addEventListener('mousemove',function(e){{
    if(!dragging)return;
    var w=Math.min(window.innerWidth*0.95,Math.max(340,window.innerWidth-e.clientX));
    panel.style.width=w+'px';
  }});
  window.addEventListener('mouseup',function(){{
    if(!dragging)return;
    dragging=false;panel.classList.remove('dragging');handle.classList.remove('dragging');
    localStorage.setItem('ci-manual-spwidth',parseInt(panel.style.width));
  }});
}})();
var SWAL_BASE={{background:undefined,customClass:{{popup:'ci-swal',confirmButton:'ci-swal-btn',cancelButton:'ci-swal-btn-cancel'}},buttonsStyling:false}};
function ciError(msg){{return Swal.fire(Object.assign({{}},SWAL_BASE,{{icon:'error',html:String(msg&&msg.message||msg),confirmButtonText:'Tamam'}}));}}
function ciPrompt(title,defaultValue){{
  return Swal.fire(Object.assign({{}},SWAL_BASE,{{icon:'question',title:title,input:'text',inputValue:defaultValue||'',
    showCancelButton:true,confirmButtonText:'Tamam',cancelButtonText:'Vazgeç'}})).then(function(r){{return r.isConfirmed?r.value:null;}});
}}
function ciToast(msg,icon){{
  return Swal.mixin({{toast:true,position:'top-end',showConfirmButton:false,timer:3200,timerProgressBar:true,
    customClass:{{popup:'ci-swal-toast'}},didOpen:function(t){{t.addEventListener('mouseenter',Swal.stopTimer);t.addEventListener('mouseleave',Swal.resumeTimer);}}
  }}).fire({{icon:icon||'success',title:msg}});
}}
function filterNav(q){{q=q.toLowerCase();document.querySelectorAll('nav .chapter').forEach(function(ch){{
  var any=false;ch.querySelectorAll('a.sec').forEach(function(a){{var m=a.textContent.toLowerCase().includes(q);a.style.display=m?'':'none';if(m)any=true;}});
  ch.style.display=(!q||any)?'':'none';}});}}
// Sıra B (kullanıcı): sol menüde Dosyalar/Sınıflar sekmesi — seçim oturumlar
// arası hatırlanır (yalnız class_tree'si olan koleksiyonlarda gösteriliyor).
function manualNavTab(btn){{
  var which=btn.dataset.tab;
  document.querySelectorAll('.navtabs button').forEach(function(b){{b.classList.remove('cur');}});
  btn.classList.add('cur');
  document.querySelectorAll('.navview').forEach(function(v){{v.classList.remove('cur');}});
  var v=document.getElementById('navview-'+which);
  if(v)v.classList.add('cur');
  localStorage.setItem('ci-manual-navtab',which);
}}
(function(){{
  var saved=localStorage.getItem('ci-manual-navtab');
  if(saved&&saved!=='files'){{
    var btn=document.querySelector('.navtabs button[data-tab="'+saved+'"]');
    if(btn)manualNavTab(btn);
  }}
}})();
async function manualAddLang(ev){{
  ev.preventDefault();
  var name=await ciPrompt('Hangi dile çevrilsin? (ör. Türkçe, Deutsch, Français)');
  if(!name||!name.trim())return;
  var code=name.trim().toLowerCase().replace(/[^a-z0-9]+/g,'-');
  var box=document.getElementById('langjob');
  box.style.display='block';box.textContent='Çeviri başlatılıyor…';
  try{{
    var r=await(await fetch('/api/manual/translate',{{method:'POST',headers:{{'Content-Type':'application/json'}},
      body:JSON.stringify({{collection:MANUAL_COLLECTION,lang:code,label:name.trim()}})}})).json();
    if(r.error){{box.textContent='❌ '+r.error;ciError(r.error);return;}}
    var t=setInterval(async function(){{
      var s=await(await fetch('/api/manual/translate-status')).json();
      if(!s.phase||s.phase==='idle')return;
      if(s.phase==='translating')box.textContent='Çevriliyor… '+(s.done||0)+'/'+(s.total||0);
      else if(s.phase==='done'){{clearInterval(t);box.textContent='✅ Hazır, yenileniyor…';ciToast('Çeviri hazır');
        location.href='/manual/'+MANUAL_COLLECTION+(MANUAL_PAGE!=='index'?'/'+MANUAL_PAGE+'.html':'')+'?lang='+code;}}
      else if(s.phase==='error'){{clearInterval(t);box.textContent='❌ '+(s.error||'');ciError(s.error||'Çeviri hatası');}}
    }},2000);
  }}catch(e){{box.textContent='❌ '+e;ciError(e);}}
}}
async function manualOpen(collection,id,btn){{
  var old=btn.textContent;btn.disabled=true;btn.textContent='…';
  try{{
    var r=await(await fetch('/api/reveal',{{method:'POST',headers:{{'Content-Type':'application/json'}},body:JSON.stringify({{collection:collection,id:id,mode:'file'}})}})).json();
    if(r.error)ciError(r.error);
  }}catch(e){{ciError(e);}}
  btn.disabled=false;btn.textContent=old;
}}
async function manualOpenFolder(collection,id,btn){{
  var old=btn.textContent;btn.disabled=true;btn.textContent='…';
  try{{
    var r=await(await fetch('/api/reveal',{{method:'POST',headers:{{'Content-Type':'application/json'}},body:JSON.stringify({{collection:collection,id:id,mode:'folder'}})}})).json();
    if(r.error)ciError(r.error);
  }}catch(e){{ciError(e);}}
  btn.disabled=false;btn.textContent=old;
}}
async function manualShowInBrowser(collection,id,secTitle,btn,hljsLang){{
  var old=btn.textContent;btn.disabled=true;btn.innerHTML='<span class="spin"></span>';
  var esc=function(s){{return String(s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');}};
  try{{
    var r=await(await fetch('/api/reveal',{{method:'POST',headers:{{'Content-Type':'application/json'}},body:JSON.stringify({{collection:collection,id:id,mode:'browser'}})}})).json();
    if(r.error){{
      document.getElementById('sp-fname').textContent='Hata';
      document.getElementById('sp-body').innerHTML='<div class="sp-err">⚠️ '+esc(r.error)+'</div>';
      SP_CONTENT='';
    }}else{{
      document.getElementById('sp-fname').textContent=r.name||r.path||secTitle;
      document.getElementById('sp-body').innerHTML='<pre><code></code></pre>';
      document.querySelector('#sp-body code').textContent=r.content;
      manualHL(document.querySelector('#sp-body code'),hljsLang);
      SP_CONTENT=r.content;
    }}
    document.getElementById('sidepanel').classList.add('show');
  }}catch(e){{ciError(e);}}
  btn.disabled=false;btn.textContent=old;
}}
// madde 2. tur, 5 (kullanıcı): body_md içindeki kod örnekleri (ör. "Public API")
// de highlight.js kullanmalı — sağ panel dışındaki TEK statik kaynak bu.
document.querySelectorAll('main pre code[class^="language-"]').forEach(function(el){{
  manualHL(el, el.className.replace('language-',''));
}});
</script></body></html>"""


def _esc(s):
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _read_source_for_static(src_root: str, unit: str) -> str | None:
    """`/api/reveal` mode=browser ile AYNI okuma mantığı — ama burada canlı
    fetch YOK, ZIP üretim ANINDA diskten okunup sayfaya GÖMÜLÜYOR (Sıra 8'in
    kısıtı: statik paket sunucu olmadan file://'dan gezilebilmeli)."""
    try:
        p = pathlib.Path(src_root) / unit
        if not p.exists() or not p.is_file() or p.stat().st_size > 3_000_000:
            return None
        return p.read_text(encoding="utf-8", errors="replace")
    except OSError:
        return None


def _static_page(model: dict, page: str, lang: str, src_root: str | None) -> str:
    """render_manual_html'in statik/taşınabilir kardeşi (Sıra 8, kullanıcı):
    TÜM hrefler GÖRECELİ (tek düz klasörde index.html + {chapter}.html yan
    yana) — dil seçici/"+ Dil ekle" ve Aç/Klasörde Aç YOK (biri canlı Ollama
    çağrısı, diğer ikisi BU makinenin dosya sistemine bağlı — taşınan ZIP'te
    hiçbiri çalışmaz). "Kaynağı Göster" kalıyor ama /api/reveal'a fetch ATMAZ:
    kaynak metni üretim anında okunup sayfaya doğrudan GÖMÜLÜ (offline çalışır)."""
    nav_html = ([f'<a class="home" href="index.html">← İndeks / Kapak</a>'] if page != "index" else []) + [
                f'<h1><a href="index.html">{_esc(model["title"])}</a></h1>',
                f'<div class="meta">{_esc(model.get("version",""))} '
                f'{"· " + _esc(model["owner"]) if model.get("owner") else ""}</div>']

    files_view = []
    for ch in model["chapters"]:
        active_ch = " active" if page == ch["slug"] else ""
        files_view.append(f'<div class="chapter"><a href="{ch["slug"]}.html" class="{active_ch}">{_esc(ch["title"])}</a>')
        for sec in ch["sections"]:
            files_view.append(f'<a class="sec" href="{ch["slug"]}.html#{sec["slug"]}">{_esc(sec["title"])}</a>')
        files_view.append("</div>")

    # Sıra B (kullanıcı): statik pakette de aynı Dosyalar/Sınıflar sekmesi —
    # veri zaten model içinde (build_manual'de bir kez hesaplandı), burada
    # yalnız GÖRECELİ hrefler ile render ediliyor.
    class_view = []
    has_any_classes = False
    for ch in model["chapters"]:
        ct = ch.get("class_tree") or {"roots": [], "other_files": []}
        if not ct["roots"] and not ct["other_files"]:
            continue
        has_any_classes = has_any_classes or bool(ct["roots"])
        active_ch = " active" if page == ch["slug"] else ""
        class_view.append(f'<div class="chapter"><a href="{ch["slug"]}.html" class="{active_ch}">{_esc(ch["title"])}</a>')
        if ct["roots"]:
            items = "".join(f'<li>{_render_ctree_node(n, ch["slug"], "", "", page == ch["slug"])}</li>' for n in ct["roots"])
            class_view.append(f'<ul class="ctree">{items}</ul>')
        if ct["other_files"]:
            of = "".join(f'<li><a href="{ch["slug"]}.html{f["href"]}">{_esc(f["title"])}</a></li>' for f in ct["other_files"])
            class_view.append(f'<div class="other"><div class="other-label">Diğer / Global</div><ul class="ctree">{of}</ul></div>')
        class_view.append("</div>")

    if has_any_classes:
        nav_html.append('<div class="navtabs">'
            '<button type="button" class="cur" data-tab="files" onclick="manualNavTab(this)">Dosyalar</button>'
            '<button type="button" data-tab="classes" onclick="manualNavTab(this)">Sınıflar</button>'
            '</div>')
        nav_html.append('<input placeholder="Ara…" onkeyup="filterNav(this.value)" id="navsearch">')
        nav_html.append(f'<div class="navview cur" id="navview-files">{"".join(files_view)}</div>')
        nav_html.append(f'<div class="navview" id="navview-classes">{"".join(class_view)}</div>')
    else:
        nav_html.append('<input placeholder="Ara…" onkeyup="filterNav(this.value)" id="navsearch">')
        nav_html.extend(files_view)
    nav = "\n".join(nav_html)

    if page == "index":
        pl = ", ".join(f"{k} ({v})" for k, v in sorted(model["stats"]["languages"].items(), key=lambda kv: -kv[1]))
        body = f"""<h1>{_esc(model["title"])}</h1>
<p>{_esc(model.get("group","") or "")}</p>
<div class="overview-grid">
  <div class="stat"><b>{model["stats"]["units"]}</b><span>Dosya</span></div>
  <div class="stat"><b>{model["stats"]["chapters"]}</b><span>Bölüm</span></div>
</div>
<p><b>Diller:</b> {_esc(pl)}</p>
{"<p><b>Kaynak:</b> " + _esc(model.get("kaynak","")) + (" — " + _esc(model.get("url","")) if model.get("url") else "") + "</p>" if model.get("kaynak") else ""}
<h2>İçindekiler</h2>
<ul>{"".join(f'<li><a href="{c["slug"]}.html">{_esc(c["title"])}</a> ({len(c["sections"])})</li>' for c in model["chapters"])}</ul>"""
    else:
        ch = next((c for c in model["chapters"] if c["slug"] == page), None)
        if ch is None:
            return "<h1>404</h1>"
        type_home_rel = model.get("type_home", {})   # zaten göreli ("chapter.html#section")
        parts = [f'<h1>{_esc(ch["title"])}</h1>']
        for sec in ch["sections"]:
            linked = _cross_link(sec["body_md"], type_home_rel, f'{ch["slug"]}.html#{sec["slug"]}')
            actions = ""
            code_html = ""
            src_text = _read_source_for_static(src_root, sec["unit"]) if src_root and sec.get("chunk_id") is not None else None
            if src_text is not None:
                # madde 2. tur, 2 (kullanıcı): canlı sistemdeki gibi SAĞ PANELDE açılsın
                # (aşağıdaki #sidepanel) — fetch YOK, kaynak <template> içine üretim
                # anında GÖMÜLÜ, panel onu oradan okuyup gösteriyor (offline çalışır).
                actions = (f'<span class="sec-actions">'
                          f'<button type="button" onclick="manualShowSource(\'code-{sec["slug"]}\',\'{_esc(sec["title"])}\',\'{_hljs_lang(sec["lang"])}\')">🖥️ Kaynağı Göster</button>'
                          f'</span>')
                code_html = f'<template id="code-{sec["slug"]}">{_esc(src_text)}</template>'
            parts.append(f'<div class="section" id="{sec["slug"]}"><h2>{_esc(sec["title"])}'
                         f'<span class="badge">{_esc(sec["lang"])}</span>{actions}</h2>'
                         f'{_md_to_html_fragment(linked, _hljs_lang(sec["lang"]))}{code_html}</div>')
        body = "\n".join(parts)

    html_lang = lang if lang in ("tr", "en") else "en"
    return f"""<!doctype html><html lang="{html_lang}"><head><meta charset="utf-8">
<title>{_esc(model["title"])}</title><style>{MANUAL_CSS}</style>
<script src="highlightjs/highlight.min.js"></script></head>
<body><nav id="manualnav">{nav}<div id="navresize"></div></nav><main>{body}</main>
<div id="sidepanel">
  <div id="sp-resize"></div>
  <div class="sp-head">
    <span class="sp-fname" id="sp-fname">…</span>
    <button class="sp-close" onclick="copySidePanel()" id="sp-copy" title="Kopyala">📋</button>
    <button class="sp-close" onclick="closeSidePanel()">✕</button>
  </div>
  <div class="sp-body" id="sp-body"></div>
</div>
<script>
function filterNav(q){{q=q.toLowerCase();document.querySelectorAll('nav .chapter').forEach(function(ch){{
  var any=false;ch.querySelectorAll('a.sec').forEach(function(a){{var m=a.textContent.toLowerCase().includes(q);a.style.display=m?'':'none';if(m)any=true;}});
  ch.style.display=(!q||any)?'':'none';}});}}
function manualNavTab(btn){{
  var which=btn.dataset.tab;
  document.querySelectorAll('.navtabs button').forEach(function(b){{b.classList.remove('cur');}});
  btn.classList.add('cur');
  document.querySelectorAll('.navview').forEach(function(v){{v.classList.remove('cur');}});
  var v=document.getElementById('navview-'+which);
  if(v)v.classList.add('cur');
  localStorage.setItem('ci-manual-navtab',which);
}}
(function(){{
  var saved=localStorage.getItem('ci-manual-navtab');
  if(saved&&saved!=='files'){{
    var btn=document.querySelector('.navtabs button[data-tab="'+saved+'"]');
    if(btn)manualNavTab(btn);
  }}
}})();
(function(){{
  var nav=document.getElementById('manualnav'),handle=document.getElementById('navresize');
  var saved=parseInt(localStorage.getItem('ci-manual-navw'));
  if(saved)nav.style.width=saved+'px';
  var dragging=false;
  handle.addEventListener('mousedown',function(e){{dragging=true;nav.classList.add('dragging');handle.classList.add('dragging');e.preventDefault();}});
  window.addEventListener('mousemove',function(e){{
    if(!dragging)return;
    var w=Math.min(window.innerWidth*0.6,Math.max(200,e.clientX));
    nav.style.width=w+'px';
  }});
  window.addEventListener('mouseup',function(){{
    if(!dragging)return;
    dragging=false;nav.classList.remove('dragging');handle.classList.remove('dragging');
    localStorage.setItem('ci-manual-navw',parseInt(nav.style.width));
  }});
}})();
var SP_CONTENT='';
var HLJS_LOADED={{}};
function manualHL(codeEl,hljsLang){{
  if(!hljsLang||typeof hljs==='undefined')return;
  codeEl.className='language-'+hljsLang;
  if(HLJS_LOADED[hljsLang]==='ok'){{hljs.highlightElement(codeEl);return;}}
  if(HLJS_LOADED[hljsLang]==='error')return;
  var s=document.createElement('script');
  s.src='highlightjs/languages/'+hljsLang+'.min.js';
  s.onload=function(){{HLJS_LOADED[hljsLang]='ok';hljs.highlightElement(codeEl);}};
  s.onerror=function(){{HLJS_LOADED[hljsLang]='error';}};
  document.head.appendChild(s);
}}
function manualShowSource(tplId,title,hljsLang){{
  var t=document.getElementById(tplId);
  if(!t)return;
  var text=t.content.textContent;
  document.getElementById('sp-fname').textContent=title;
  document.getElementById('sp-body').innerHTML='<pre><code></code></pre>';
  document.querySelector('#sp-body code').textContent=text;
  manualHL(document.querySelector('#sp-body code'),hljsLang);
  SP_CONTENT=text;
  document.getElementById('sidepanel').classList.add('show');
}}
function closeSidePanel(){{document.getElementById('sidepanel').classList.remove('show');}}
async function copySidePanel(){{
  if(!SP_CONTENT)return;
  try{{
    await navigator.clipboard.writeText(SP_CONTENT);
    var b=document.getElementById('sp-copy'),old=b.textContent;b.textContent='✅';
    setTimeout(function(){{b.textContent=old;}},1200);
  }}catch(e){{}}
}}
document.addEventListener('keydown',function(e){{
  if(e.key==='Escape'&&document.getElementById('sidepanel').classList.contains('show'))closeSidePanel();
}});
(function(){{
  var panel=document.getElementById('sidepanel'),handle=document.getElementById('sp-resize');
  var saved=parseInt(localStorage.getItem('ci-manual-spwidth'));
  if(saved)panel.style.width=saved+'px';
  var dragging=false;
  handle.addEventListener('mousedown',function(e){{dragging=true;panel.classList.add('dragging');handle.classList.add('dragging');e.preventDefault();}});
  window.addEventListener('mousemove',function(e){{
    if(!dragging)return;
    var w=Math.min(window.innerWidth*0.95,Math.max(340,window.innerWidth-e.clientX));
    panel.style.width=w+'px';
  }});
  window.addEventListener('mouseup',function(){{
    if(!dragging)return;
    dragging=false;panel.classList.remove('dragging');handle.classList.remove('dragging');
    localStorage.setItem('ci-manual-spwidth',parseInt(panel.style.width));
  }});
}})();
document.querySelectorAll('main pre code[class^="language-"]').forEach(function(el){{
  manualHL(el, el.className.replace('language-',''));
}});
</script></body></html>"""


def render_manual_zip(model: dict, lang: str = "") -> bytes:
    """Sıra 8 (kullanıcı): "statik html export eklenmesini istiyorum. Bütün
    dosyalar tek zip dosyasında." `model` ÇAĞIRAN tarafından zaten
    model_for_lang() ile dönüştürülmüş olmalı (docx/pdf render'larıyla AYNI
    sözleşme) — burada yalnız `lang` html lang= özniteliği için kullanılır.

    Sıra 3/4/5 (kullanıcı, highlight.js): statik paket `/static/vendor/...`
    (SUNUCU-bağımlı, mutlak yol) kullanamaz — bu yüzden highlight.js çekirdeği
    + YALNIZ bu manuelde GERÇEKTEN kullanılan dil dosyaları ZIP'in İÇİNE,
    GÖRECELİ "highlightjs/" alt klasörü olarak kopyalanır (192 dilin TAMAMI
    değil — çoğu koleksiyon 1-3 dil kullanır, gereksiz şişirme önlenir)."""
    active_lang = lang or model.get("lang", "en")
    prof = retrieval.get_profile_payload(model["collection"])
    src_root = prof.get("path") or None
    needed_hljs = sorted({hl for ch in model["chapters"] for sec in ch["sections"]
                          if (hl := _hljs_lang(sec.get("lang", "")))})

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("index.html", _static_page(model, "index", active_lang, src_root))
        for ch in model["chapters"]:
            zf.writestr(f'{ch["slug"]}.html', _static_page(model, ch["slug"], active_lang, src_root))
        hljs_core = ROOT / "static" / "vendor" / "highlightjs" / "highlight.min.js"
        if hljs_core.exists():
            zf.write(hljs_core, "highlightjs/highlight.min.js")
            for hl in needed_hljs:
                langfile = ROOT / "static" / "vendor" / "highlightjs" / "languages" / f"{hl}.min.js"
                if langfile.exists():
                    zf.write(langfile, f"highlightjs/languages/{hl}.min.js")
    return buf.getvalue()


# ---------------- DOCX / PDF (aynı model, farklı render — HTML'in aksine
# çapraz-link YOK; sade akış metni, kod blokları anlaşılır biçimde ayrılmış) ----------------

# Sıra 3. tur (kullanıcı, "PDF/DOCX'te de syntax highlighting kullanılacak" —
# önceden kapsam dışı bırakılmıştı, kullanıcı geri açtı): highlight.js
# TARAYICI'da çalışır, DOCX/PDF'in DOM'u yok — bu yüzden burada AYRI, Python
# tarafında çalışan bir kütüphane (pygments) kullanılıyor. python-docx/
# reportlab zaten LAZY import ediliyor (bkz. render_manual_docx/pdf), pygments
# de AYNI disiplinle burada lazy import edilir.
def _pygments_color_for(token_type) -> str | None:
    """3 renkli kısıtlı eşleme — HTML/CSS temasının AÇIK-ZEMİN ("-d", koyu)
    varyantlarıyla tutarlı (DOCX/PDF sayfaları beyaz zemin, panelin koyu
    zemini DEĞİL): Keyword/Number->amber-d, String/Name.Class/Function->
    teal-d, Comment->dim gri. Geri kalanı (Punctuation/Text/Operator vb.)
    None döner — sayfanın varsayılan (siyah) metin rengiyle kalır."""
    from pygments.token import Token
    if token_type in Token.Keyword or token_type in Token.Operator.Word or token_type in Token.Literal.Number:
        return "C58A37"
    if (token_type in Token.Literal.String or token_type in Token.Name.Class
            or token_type in Token.Name.Function or token_type in Token.Name.Builtin):
        return "2F7E6C"
    if token_type in Token.Comment:
        return "6F6555"
    return None


def _pygments_lex(code: str, lang: str) -> list[tuple[str | None, str]]:
    """(hex_renk_veya_None, metin_parçası) çiftleri — DOCX (per-run renk) ve
    PDF (reportlab <font color> markup) render'ları AYNI tek listeyi kullanır
    (HTML/DOCX/PDF'in "tek model, üç render" ilkesiyle tutarlı). `lang` boşsa
    veya pygments o dil için bir lexer bulamazsa (ör. proje-özel bir dil adı)
    SESSİZCE tek parça, renksiz metne düşer — üretim asla çökmez."""
    if not lang:
        return [(None, code)]
    try:
        from pygments import lex
        from pygments.lexers import get_lexer_by_name
        lexer = get_lexer_by_name(lang, stripnl=False)
        return [(_pygments_color_for(tok), val) for tok, val in lex(code, lexer)]
    except Exception:
        return [(None, code)]


def _plain_blocks(body_md: str) -> list[tuple[str, str]]:
    """(tür, metin) çiftleri: tür = "h2"|"h3"|"p"|"li"|"code". Basit satır bazlı
    ayrıştırma — hem DOCX hem PDF render'ı bunu ortak kullanır."""
    out = []
    in_code = False
    code_buf = []
    for line in body_md.split("\n"):
        if line.strip().startswith("```"):
            if in_code:
                out.append(("code", "\n".join(code_buf))); code_buf = []
            in_code = not in_code
            continue
        if in_code:
            code_buf.append(line); continue
        m = re.match(r"^(#{1,3})\s+(.*)", line)
        if m:
            out.append((f"h{len(m.group(1))}", m.group(2).strip())); continue
        m = re.match(r"^[-*]\s+(.*)", line)
        if m:
            out.append(("li", m.group(1).strip())); continue
        if line.strip():
            out.append(("p", re.sub(r"[*`]", "", line.strip())))
    return out


def render_manual_docx(model: dict) -> bytes:
    """python-docx ile — bkz. skill notları: sayfa/tablo/liste tuzakları
    burada yok (yalnız başlık+paragraf+kod-paragrafı kullanılıyor, düz akış)."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading(model["title"], level=0)
    meta = f'{model.get("version","")}  {model.get("owner","")}  {model.get("group","")}'.strip()
    if meta:
        p = doc.add_paragraph(meta); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'{model["stats"]["units"]} dosya · {model["stats"]["chapters"]} bölüm — '
                      f'{datetime.now(timezone.utc).strftime("%Y-%m-%d")}')

    for ch in model["chapters"]:
        doc.add_heading(ch["title"], level=1)
        for sec in ch["sections"]:
            doc.add_heading(f'{sec["title"]}  [{sec["lang"]}]', level=2)
            for kind, text in _plain_blocks(sec["body_md"]):
                if kind in ("h2", "h3"):
                    doc.add_heading(text, level=3)
                elif kind == "li":
                    doc.add_paragraph(text, style="List Bullet")
                elif kind == "code":
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Pt(18)
                    for color, tok_text in _pygments_lex(text, sec["lang"]):
                        r = p.add_run(tok_text); r.font.name = "Consolas"; r.font.size = Pt(9)
                        if color:
                            r.font.color.rgb = RGBColor.from_string(color)
                else:
                    doc.add_paragraph(text)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# reportlab'ın yerleşik Helvetica/Times fontları Latin-1/WinAnsi ile sınırlı —
# Türkçe'ye özgü ı/ş/ğ glifleri YOK (canlı testte "Kullanım" -> "Kullan■m" olarak
# bozulduğu doğrulandı). Windows'ta zaten kurulu gerçek Unicode TTF'ler
# (Georgia/Calibri/Consolas — HTML tasarımıyla aynı aile) kaydedilip kullanılır.
_FONTS_REGISTERED = False
_FONT_DIR = pathlib.Path(r"C:\Windows\Fonts")
_FONT_FILES = {"Georgia": "georgia.ttf", "Georgia-Bold": "georgiab.ttf",
               "Calibri": "calibri.ttf", "Calibri-Bold": "calibrib.ttf",
               "Consolas": "consola.ttf"}

def _ensure_pdf_fonts():
    """Idempotent — Windows dışı bir ortamda TTF bulunamazsa yerleşik fontlara
    sessizce düşülür (Türkçe glif eksik kalır ama üretim çökmez)."""
    global _FONTS_REGISTERED
    if _FONTS_REGISTERED:
        return True
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    try:
        for name, fname in _FONT_FILES.items():
            pdfmetrics.registerFont(TTFont(name, str(_FONT_DIR / fname)))
        _FONTS_REGISTERED = True
    except Exception:
        pass
    return _FONTS_REGISTERED


def render_manual_pdf(model: dict) -> bytes:
    """reportlab Platypus ile — pypdf/weasyprint/playwright DEĞİL: saf Python,
    yeni sistem bağımlılığı yok (Windows'ta Cairo/Pango riski, Chromium indirmesi
    yok) — pdf skill'inin önerdiği CREATE yolu."""
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, ListFlowable, ListItem

    have_unicode_fonts = _ensure_pdf_fonts()
    f_head = "Georgia" if have_unicode_fonts else "Helvetica-Bold"
    f_body = "Calibri" if have_unicode_fonts else "Helvetica"
    f_code = "Consolas" if have_unicode_fonts else "Courier"

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2.2 * cm, bottomMargin=2 * cm,
                            leftMargin=2.2 * cm, rightMargin=2.2 * cm)
    styles = getSampleStyleSheet()
    amber, teal, dim = colors.HexColor("#c58a37"), colors.HexColor("#2f7e6c"), colors.HexColor("#6f6555")
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName=f_head, textColor=amber, spaceAfter=10)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName=f_head, textColor=amber, spaceBefore=14, spaceAfter=6)
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontName=f_head, textColor=teal, spaceBefore=8, spaceAfter=4)
    body = ParagraphStyle("Body", parent=styles["Normal"], fontName=f_body, fontSize=10, leading=14, spaceAfter=6)
    codep = ParagraphStyle("Code", parent=styles["Code"], fontName=f_code, fontSize=8, leading=11,
                           backColor=colors.HexColor("#f2f0ea"), leftIndent=10, spaceAfter=8)
    metap = ParagraphStyle("Meta", parent=styles["Normal"], fontName=f_body, fontSize=9, textColor=dim, spaceAfter=16)

    def esc(s): return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story = [Paragraph(esc(model["title"]), h1)]
    meta = " · ".join(x for x in (model.get("version"), model.get("owner"), model.get("group")) if x)
    if meta:
        story.append(Paragraph(esc(meta), metap))
    story.append(Paragraph(f'{model["stats"]["units"]} dosya · {model["stats"]["chapters"]} bölüm', metap))
    story.append(PageBreak())

    for ch in model["chapters"]:
        story.append(Paragraph(esc(ch["title"]), h1))
        for sec in ch["sections"]:
            story.append(Paragraph(f'{esc(sec["title"])} <font color="#{teal.hexval()[2:]}" size="8">[{esc(sec["lang"])}]</font>', h2))
            items = []
            for kind, text in _plain_blocks(sec["body_md"]):
                if kind in ("h2", "h3"):
                    story.append(Paragraph(esc(text), h3))
                elif kind == "li":
                    items.append(ListItem(Paragraph(esc(text), body)))
                elif kind == "code":
                    marked = "".join(
                        (f'<font color="#{color}">{esc(tok).replace(chr(10), "<br/>")}</font>' if color
                         else esc(tok).replace(chr(10), "<br/>"))
                        for color, tok in _pygments_lex(text, sec["lang"]))
                    story.append(Paragraph(marked, codep))
                else:
                    if items:
                        story.append(ListFlowable(items, bulletType="bullet")); items = []
                    story.append(Paragraph(esc(text), body))
            if items:
                story.append(ListFlowable(items, bulletType="bullet"))
        story.append(PageBreak())
    doc.build(story)
    return buf.getvalue()
